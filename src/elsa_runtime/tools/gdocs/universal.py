"""
Universal Drive document reader.

Handles the Anthropic-managed gdoc_read limitation: it only accepts native
Google Docs (`application/vnd.google-apps.document`). Drive often contains
Word .docx, .pdf, plain .txt, etc. — uploaded files keep their native format
even when shown alongside Google Docs in Drive UI.

This module:
1. Looks up Drive metadata to detect mime_type.
2. Routes to:
   - Google Docs API for native Docs (delegates to GoogleDocsReader)
   - Drive download + python-docx for .docx
   - Drive download + pymupdf for .pdf
   - Drive download + utf-8 decode for .txt / .md
3. For non-native types, also persists the downloaded blob to a temp dir
   so the user can open / re-process the original.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from .reader import GoogleDocsReader

# MIME type constants
MIME_GOOGLE_DOC = "application/vnd.google-apps.document"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_DOC_LEGACY = "application/msword"
MIME_PDF = "application/pdf"
MIME_TXT = "text/plain"
MIME_MARKDOWN = "text/markdown"

# Default save location for downloaded blobs (matches gmail attachments pattern).
DEFAULT_TEMP_DIR = Path.home() / "Projects/elsa-data/temp/gdocs"


class UniversalDocReader:
    """Read any Drive doc — native Google Doc, .docx, .pdf, .txt, .md.

    Constructor takes both Drive and Docs services because routing depends
    on metadata from Drive but native Docs use the Docs API.
    """

    def __init__(
        self,
        drive_service: Any,
        docs_service: Any,
        temp_dir: Path | str = DEFAULT_TEMP_DIR,
    ) -> None:
        self._drive = drive_service
        self._docs_reader = GoogleDocsReader(docs_service)
        self._temp_dir = Path(temp_dir).expanduser()

    def read(self, file_id: str) -> dict:
        """Read a Drive item, dispatching by mime_type.

        Returns:
            {
              "id": str,
              "name": str,
              "mime_type": str,
              "method": "google_docs" | "docx" | "pdf" | "text",
              "text": str,                  # extracted plain text
              "headings": [{"level": int, "text": str}],  # if applicable
              "char_count": int,
              "saved_to": str | None,        # local path for non-native types
            }
        """
        meta = (
            self._drive.files()
            .get(
                fileId=file_id,
                fields="id, name, mimeType, size",
            )
            .execute()
        )
        mime = meta.get("mimeType", "")
        name = meta.get("name", "")

        if mime == MIME_GOOGLE_DOC:
            return self._read_google_doc(file_id, name)
        elif mime in (MIME_DOCX, MIME_DOC_LEGACY):
            return self._read_docx(file_id, name, mime)
        elif mime == MIME_PDF:
            return self._read_pdf(file_id, name)
        elif mime in (MIME_TXT, MIME_MARKDOWN) or mime.startswith("text/"):
            return self._read_text(file_id, name, mime)
        else:
            # Unknown mime — download anyway, give path back for user inspection.
            return self._download_unknown(file_id, name, mime)

    # ── Native Google Doc ─────────────────────────────────────────────

    def _read_google_doc(self, file_id: str, name: str) -> dict:
        result = self._docs_reader.read(file_id)
        return {
            "id": file_id,
            "name": result.get("title") or name,
            "mime_type": MIME_GOOGLE_DOC,
            "method": "google_docs",
            "text": result["text"],
            "headings": result["headings"],
            "char_count": result["char_count"],
            "saved_to": None,
        }

    # ── DOCX ──────────────────────────────────────────────────────────

    def _read_docx(self, file_id: str, name: str, mime: str) -> dict:
        blob = self._download_media(file_id)
        save_path = self._save_blob(file_id, name or f"{file_id}.docx", blob)

        # Parse with python-docx
        import docx as docx_lib  # local import keeps top-level light

        doc = docx_lib.Document(io.BytesIO(blob))
        text_parts: list[str] = []
        headings: list[dict] = []

        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                continue
            text_parts.append(text)

            # python-docx style names for headings: "Heading 1", "Heading 2", ...
            style_name = (para.style.name or "") if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                headings.append({"level": level, "text": text.strip()})

        # Tables: include cell text linearly so search can hit it.
        for tbl_idx, table in enumerate(doc.tables):
            text_parts.append(f"[TABLE {tbl_idx}]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        return {
            "id": file_id,
            "name": name,
            "mime_type": mime,
            "method": "docx",
            "text": full_text,
            "headings": headings,
            "char_count": len(full_text),
            "saved_to": str(save_path),
        }

    # ── PDF ───────────────────────────────────────────────────────────

    def _read_pdf(self, file_id: str, name: str) -> dict:
        blob = self._download_media(file_id)
        save_path = self._save_blob(file_id, name or f"{file_id}.pdf", blob)

        import pymupdf

        doc = pymupdf.open(stream=blob, filetype="pdf")
        text_parts: list[str] = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()

        full_text = "\n".join(text_parts)
        return {
            "id": file_id,
            "name": name,
            "mime_type": MIME_PDF,
            "method": "pdf",
            "text": full_text,
            "headings": [],  # No structural extraction for PDF here
            "char_count": len(full_text),
            "saved_to": str(save_path),
        }

    # ── Plain text / Markdown ─────────────────────────────────────────

    def _read_text(self, file_id: str, name: str, mime: str) -> dict:
        blob = self._download_media(file_id)
        ext = ".md" if mime == MIME_MARKDOWN else ".txt"
        save_path = self._save_blob(file_id, name or f"{file_id}{ext}", blob)
        text = blob.decode("utf-8", errors="replace")
        return {
            "id": file_id,
            "name": name,
            "mime_type": mime,
            "method": "text",
            "text": text,
            "headings": [],
            "char_count": len(text),
            "saved_to": str(save_path),
        }

    # ── Unknown mime: download + return path ──────────────────────────

    def _download_unknown(self, file_id: str, name: str, mime: str) -> dict:
        try:
            blob = self._download_media(file_id)
        except Exception as e:
            return {
                "id": file_id,
                "name": name,
                "mime_type": mime,
                "method": "unsupported",
                "text": "",
                "headings": [],
                "char_count": 0,
                "saved_to": None,
                "error": f"Cannot download (mime={mime}): {e}",
            }
        save_path = self._save_blob(file_id, name or file_id, blob)
        return {
            "id": file_id,
            "name": name,
            "mime_type": mime,
            "method": "downloaded_only",
            "text": "",
            "headings": [],
            "char_count": 0,
            "saved_to": str(save_path),
            "note": (
                f"Mime '{mime}' not directly extractable; downloaded blob "
                f"saved for manual inspection."
            ),
        }

    # ── Helpers ───────────────────────────────────────────────────────

    def _download_media(self, file_id: str) -> bytes:
        """Download Drive file bytes via files.get_media (alt=media)."""
        from googleapiclient.http import MediaIoBaseDownload

        request = self._drive.files().get_media(fileId=file_id)
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _status, done = downloader.next_chunk()
        return buf.getvalue()

    def _save_blob(self, file_id: str, filename: str, blob: bytes) -> Path:
        """Write blob to temp_dir/<file_id>/<filename>. Returns path."""
        # Sanitize filename (basic): replace path separators
        safe_name = filename.replace("/", "_").replace("\\", "_")
        dest_dir = self._temp_dir / file_id
        dest_dir.mkdir(parents=True, exist_ok=True)
        path = dest_dir / safe_name
        path.write_bytes(blob)
        return path
